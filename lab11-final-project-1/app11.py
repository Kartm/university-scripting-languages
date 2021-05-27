import io
import math
import re
from io import BytesIO

import matplotlib.pyplot as plt
import numpy
import requests
from PIL import Image
import docx


# book downloaded from https://www.gutenberg.org/files/52466/52466-h/52466-h.htm

class ChapterStats:
    def __init__(self, paragraphs):
        self.paragraphs_count = len(paragraphs)
        self.words_count = sum([len(p.split()) for p in paragraphs])
        self.min_words = min([len(p.split()) for p in paragraphs])
        self.max_words = max([len(p.split()) for p in paragraphs])
        self.avg_words = self.words_count / self.paragraphs_count


def read_file(path: str):
    with open(path) as f:
        content = f.read()
        return content


def parse_book(plain_text: str):
    title = re.search(r'Title: (?P<title>.*)', plain_text).group('title')
    author = re.search(r'Author: (?P<author>.*)', plain_text).group('author')
    chapter_lines = plain_text.split("\n")[870:1166]

    return title, author, chapter_lines


def get_paragraphs_from_lines(chapter_lines):
    paragraphs = list()
    current_paragraph = ''
    for line in chapter_lines:
        stripped_line = line.strip()

        current_paragraph += ' ' + stripped_line if stripped_line else ''
        # print(current_paragraph)
        if stripped_line == '' and current_paragraph != '':
            paragraphs.append(current_paragraph)
            current_paragraph = ''

    return paragraphs


def get_paragraph_data(paragraphs) -> (BytesIO, ChapterStats):
    def round_to_10(x):
        return int(math.ceil(x / 10.0)) * 10

    paragraphs_word_count = [round_to_10(len(p.split())) for p in paragraphs]
    print(f'Sorted paragraph count: {sorted(paragraphs_word_count)}')
    duplicate_count = {e: paragraphs_word_count.count(e) for e in paragraphs_word_count}
    print(f'Duplicate count: {duplicate_count}')

    plt.hist(paragraphs_word_count, bins=numpy.arange(0, 530, 10))
    plt.xlabel('No. of words')
    plt.ylabel('No. of paragraphs')
    plt.title('Chapter I - paragraph word count distribution')

    stream = io.BytesIO()
    plt.savefig(stream, format='png')
    stream.seek(0)

    stats = ChapterStats(paragraphs)

    return stream, stats


def download_image() -> Image:
    url = 'https://www.gutenberg.org/files/52466/52466-h/images/i_005.jpg'
    original_image = Image.open(requests.get(url, stream=True).raw)
    original_image = original_image.convert("RGB")

    smaller_edge = min(original_image.width, original_image.height)
    square_image = original_image.crop((0, 0, smaller_edge, smaller_edge))

    resized_image = square_image.resize((smaller_edge * 2, smaller_edge * 2))

    return resized_image


def load_local_image() -> Image:
    original_image = Image.open("./dinosaur.png")

    black_and_white_image = original_image.convert('LA')

    rotated_image = black_and_white_image.rotate(-25)

    return rotated_image


def get_combined_pictures_stream(bottom_image, top_image) -> BytesIO:
    combined_picture = bottom_image.copy()
    combined_picture.paste(top_image, box=(100, 300), mask=top_image.convert("RGBA"))

    stream = io.BytesIO()
    stream.seek(0)
    combined_picture.save(stream, format="PNG")
    return stream


def make_word_document(title, author, picture_stream: BytesIO, paragraph_plot_stream: BytesIO,
                       chapter_stats: ChapterStats):
    doc = docx.Document()

    doc.add_paragraph(title, 'Title')
    doc.add_picture(picture_stream)
    doc.add_paragraph(author, 'Subtitle')
    doc.add_paragraph('Report author: Łukasz Blachnicki', 'Caption').bold = True

    doc.add_page_break()

    doc.add_paragraph('Content analysis', 'Title')
    doc.add_picture(paragraph_plot_stream)

    doc.add_paragraph('The chart shows the distribution of words count in paragraphs of Chapter I. Basic metrics of '
                      'the first '
                      'chapter:')
    doc.add_paragraph(
        f'Number of paragraphs: {chapter_stats.paragraphs_count}', style='List Bullet'
    )
    doc.add_paragraph(
        f'Number of words: {chapter_stats.words_count}', style='List Bullet'
    )
    doc.add_paragraph(
        f'Minimal number of words in a paragraph: {chapter_stats.min_words}', style='List Bullet'
    )
    doc.add_paragraph(
        f'Maximal number of words in a paragraph: {chapter_stats.max_words}', style='List Bullet'
    )
    doc.add_paragraph(
        f'Average number of words in a paragraph: {chapter_stats.avg_words}', style='List Bullet'
    )

    doc.add_paragraph(
        'Book source: https://www.gutenberg.org/files/52466/52466-h/52466-h.htm'
    )

    doc.save('document.docx')


def run():
    title, author, chapter_lines = parse_book(read_file("./52466-0.txt"))
    plot_stream, stats = get_paragraph_data(get_paragraphs_from_lines(chapter_lines))
    picture_stream = get_combined_pictures_stream(download_image(), load_local_image())
    make_word_document(title, author, picture_stream, plot_stream, stats)


if __name__ == "__main__":
    run()
